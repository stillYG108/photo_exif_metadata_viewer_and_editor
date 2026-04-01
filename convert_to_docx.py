import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

def add_page_break(doc):
    """Add page break"""
    doc.add_page_break()

def set_document_margins(doc, top=1, bottom=1, left=1, right=1):
    """Set document margins (in inches)"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_header_footer(doc):
    """Add header and footer to document"""
    section = doc.sections[0]
    
    # Add footer with page number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    footer_run = footer_para.add_run("Page ")
    footer_run.font.size = Pt(10)
    
    # Add page number
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    footer_para._p.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    footer_para._p.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_para._p.append(fldChar2)

def parse_and_format_text(paragraph, text):
    """Parse markdown formatting and apply to paragraph"""
    if not text or not text.strip():
        return
    
    # Pattern to match: **bold**, __bold__, `code`, *italic*, _italic_
    pattern = r'\*\*(.+?)\*\*|__(.+?)__|`(.+?)`|\*(.+?)\*|_(.+?)_'
    
    current_pos = 0
    
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > current_pos:
            before_text = text[current_pos:match.start()]
            if before_text:
                paragraph.add_run(before_text)
        
        matched_text = (match.group(1) or match.group(2) or 
                       match.group(3) or match.group(4) or match.group(5))
        
        if match.group(1) or match.group(2):  # Bold
            run = paragraph.add_run(matched_text)
            run.bold = True
            run.font.name = 'Calibri'
        elif match.group(3):  # Code
            run = paragraph.add_run(matched_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(192, 0, 0)
        elif match.group(4) or match.group(5):  # Italic
            run = paragraph.add_run(matched_text)
            run.italic = True
            run.font.name = 'Calibri'
        
        current_pos = match.end()
    
    # Add remaining text
    if current_pos < len(text):
        remaining = text[current_pos:]
        if remaining:
            paragraph.add_run(remaining)

def read_markdown(filepath):
    """Read markdown file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def convert_markdown_to_docx_ieee(markdown_content, output_path):
    """Convert markdown content to Word document with IEEE formatting"""
    doc = Document()
    
    # Set margins
    set_document_margins(doc, top=1, bottom=1, left=1, right=1)
    
    # Add header/footer
    add_header_footer(doc)
    
    # Set default font to Calibri
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    lines = markdown_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines at start
        if not line.strip() and not in_code_block:
            i += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                i += 1
                continue
            else:
                in_code_block = False
                # Add code block to document
                if code_lines:
                    # Add code block with formatting
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    
                    # Add background styling
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E8E8E8')
                    p._element.get_or_add_pPr().append(shading_elm)
                    
                    for code_line in code_lines:
                        if code_line or code_lines.index(code_line) < len(code_lines) - 1:
                            run = p.add_run(code_line + '\n')
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                    
                    code_lines = []
                i += 1
                continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Remove leading/trailing whitespace
        line = line.strip()
        
        # Heading 1
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            # Apply IEEE font
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.bold = True
            i += 1
            continue
        
        # Heading 2
        if line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.bold = True
            i += 1
            continue
        
        # Heading 3
        if line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.bold = True
            i += 1
            continue
        
        # Heading 4
        if line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.bold = True
            i += 1
            continue
        
        # Heading 5
        if line.startswith('##### '):
            p = doc.add_heading(line[6:].strip(), level=5)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            i += 1
            continue
        
        # Horizontal divider for page break
        if line.startswith('---'):
            add_page_break(doc)
            i += 1
            continue
        
        # Bullet list
        if line.lstrip().startswith('- '):
            bullet_text = line.lstrip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            parse_and_format_text(p, bullet_text)
            for run in p.runs:
                if not run.font.name or run.font.name == 'Calibri':
                    run.font.name = 'Calibri'
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', line.lstrip()):
            match = re.match(r'^(\d+)\.\s(.+)', line.lstrip())
            if match:
                num_text = match.group(2).strip()
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                parse_and_format_text(p, num_text)
                for run in p.runs:
                    if not run.font.name or run.font.name == 'Calibri':
                        run.font.name = 'Calibri'
            i += 1
            continue
        
        # Table handling (simple)
        if '|' in line:
            # For now, just add as formatted text
            if line.count('|') > 1:
                # Skip table header separator lines
                if '---' not in line and 'Requirement' not in line and 'Format' not in line:
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if cells:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        for j, cell in enumerate(cells):
                            if j > 0:
                                p.add_run(' | ')
                            run = p.add_run(cell)
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10)
            i += 1
            continue
        
        # Regular paragraph
        if line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            parse_and_format_text(p, line)
            
            # Ensure all runs have proper font
            for run in p.runs:
                if not run.font.name or run.font.name == 'Arial':
                    run.font.name = 'Calibri'
                if not run.font.size:
                    run.font.size = Pt(11)
        
        i += 1
    
    # Save document with proper properties
    doc.save(output_path)
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = "Photo EXIF Metadata Viewer and Editor - Application Report"
    core_props.subject = "Comprehensive Application Development Report"
    core_props.author = "Development Team"
    core_props.comments = "Professional IEEE formatted application report"
    
    print(f"✅ IEEE formatted document saved: {output_path}")

if __name__ == "__main__":
    markdown_file = r"c:\Users\yugm chaudhary\Desktop\photo_exif_metadata_viewer_and_editor\APP_REPORT.md"
    output_file = r"c:\Users\yugm chaudhary\Desktop\photo_exif_metadata_viewer_and_editor\APP_REPORT.docx"
    
    print("Converting markdown to IEEE formatted Word document...")
    markdown_content = read_markdown(markdown_file)
    convert_markdown_to_docx_ieee(markdown_content, output_file)
    print("✅ Conversion complete with IEEE formatting!")
